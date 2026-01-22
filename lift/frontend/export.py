import io
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from lift.frontend.utils import get_label, get_version


class WordReport:
    def __init__(self, plots: dict, inputs: pd.DataFrame, language: dict[str, str]):
        self.plots = plots
        self.inputs = inputs
        # create report method is executed by streamlit download button.
        # Therefore, it must not contain any streamlit functionality, such as st.session_state which stores the language package.
        # As a workaround, we pass the language package directly to the WordReport class.
        self.language = language

    def create_report(
        self,
    ):
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        style = doc.styles["Caption"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(10)

        style = doc.styles["Title"]
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        style.font.color.rgb = RGBColor(0, 0, 0)

        style = doc.styles["Heading 1"]
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        style.font.color.rgb = RGBColor(0, 0, 0)

        # Add title followed by empty line
        heading = doc.add_heading(get_label("report.title", language=self.language), level=0)
        heading.alignment = 1  # Center alignment

        def format_cell(cell):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()

            return paragraph, run

        kpi_width = Inches(1.5)
        kpi_plots = self.plots["kpi_diagrams"]

        paragraph = doc.add_paragraph("")
        run = paragraph.add_run()
        run.add_break(WD_BREAK.LINE)
        run.add_break(WD_BREAK.LINE)
        kpi1 = doc.add_table(rows=2, cols=2)
        hdr_cells = kpi1.rows[0].cells
        cell_paragraph, _ = format_cell(hdr_cells[0])
        cell_paragraph.text = get_label("main.kpi_diagrams.costs.title", language=self.language)
        cell_paragraph, _ = format_cell(hdr_cells[1])
        cell_paragraph.text = get_label("main.kpi_diagrams.emissions.title", language=self.language)

        pic_cells = kpi1.rows[1].cells
        _, cell_run = format_cell(pic_cells[0])
        cell_run.add_picture(kpi_plots["kpi_diagrams.costs"].plot_bytestream, width=kpi_width)

        _, cell_run = format_cell(pic_cells[1])
        cell_run.add_picture(kpi_plots["kpi_diagrams.emissions"].plot_bytestream, width=kpi_width)

        paragraph = doc.add_paragraph("")
        run = paragraph.add_run()
        run.add_break(WD_BREAK.LINE)
        run.add_break(WD_BREAK.LINE)

        kpi2 = doc.add_table(rows=2, cols=3)
        hdr_cells = kpi2.rows[0].cells
        cell_paragraph, _ = format_cell(hdr_cells[0])
        cell_paragraph.text = get_label("main.kpi_diagrams.self_sufficiency.title", language=self.language)
        cell_paragraph, _ = format_cell(hdr_cells[1])
        cell_paragraph.text = get_label("main.kpi_diagrams.self_consumption.title", language=self.language)
        cell_paragraph, _ = format_cell(hdr_cells[2])
        cell_paragraph.text = get_label("main.kpi_diagrams.home_charging.title", language=self.language)

        pic_cells = kpi2.rows[1].cells
        _, cell_run = format_cell(pic_cells[0])
        cell_run.add_picture(kpi_plots["kpi_diagrams.self_sufficiency"].plot_bytestream, width=kpi_width)

        _, cell_run = format_cell(pic_cells[1])
        cell_run.add_picture(kpi_plots["kpi_diagrams.self_consumption"].plot_bytestream, width=kpi_width)

        _, cell_run = format_cell(pic_cells[2])
        cell_run.add_picture(kpi_plots["kpi_diagrams.home_charging"].plot_bytestream, width=kpi_width)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        time_plots = self.plots["time_diagrams"]
        doc.add_heading(get_label("main.time_diagrams.costs.title.label", language=self.language), level=1)
        paragraph = doc.add_paragraph("")
        run = paragraph.add_run()
        run.add_break(WD_BREAK.LINE)
        run.add_picture(time_plots["time_diagrams.costs"].plot_bytestream, width=Inches(6))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        doc.add_heading(get_label("main.time_diagrams.emissions.title.label", language=self.language), level=1)
        paragraph = doc.add_paragraph("")
        run = paragraph.add_run()
        run.add_break(WD_BREAK.LINE)
        run.add_picture(time_plots["time_diagrams.emissions"].plot_bytestream, width=Inches(6))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        doc.add_heading(get_label("report.table_inputs.title", language=self.language), level=1)
        table = doc.add_table(rows=len(self.inputs) + 1, cols=4)
        table.style = "Table Grid"
        table.style = "Light List"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = get_label("report.table_inputs.columns.block", language=self.language)
        hdr_cells[1].text = get_label("report.table_inputs.columns.parameter", language=self.language)
        hdr_cells[
            2
        ].text = f"{get_label('report.table_inputs.columns.value', language=self.language)} ({get_label('main.name_baseline', language=self.language)})"
        hdr_cells[
            3
        ].text = f"{get_label('report.table_inputs.columns.value', language=self.language)} ({get_label('main.name_expansion', language=self.language)})"

        tbl_header = OxmlElement("w:tblHeader")  # create new oxml element flag which indicates that row is header row
        first_row_props = table.rows[
            0
        ]._element.get_or_add_trPr()  # get if exists or create new table row properties el
        first_row_props.append(tbl_header)  # now first row is the header row

        for row_idx, (index, row) in enumerate(self.inputs.iterrows()):
            cells = table.rows[row_idx + 1].cells
            cells[0].text = index[0]
            cells[1].text = index[1]
            cells[2].text = str(row["baseline"])
            cells[3].text = str(row["expansion"])

            if row["baseline"] != row["expansion"]:
                for cell in cells:
                    # Access the cell's XML to set the background color
                    cell_xml = cell._tc

                    # Create a new XML element for the background
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "FFFF00")  # Hex code for yellow

                    # Add the shading element to the cell's XML
                    cell_xml.get_or_add_tcPr().append(shading)

        def add_page_number(run):
            def create_element(name):
                return OxmlElement(name)

            def create_attribute(element, name, value):
                element.set(ns.qn(name), value)

            fldChar1 = create_element("w:fldChar")
            create_attribute(fldChar1, "w:fldCharType", "begin")

            instrText = create_element("w:instrText")
            create_attribute(instrText, "xml:space", "preserve")
            instrText.text = "PAGE"

            fldChar2 = create_element("w:fldChar")
            create_attribute(fldChar2, "w:fldCharType", "end")

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        footer_paragraph = doc.sections[0].footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(doc.sections[0].footer.paragraphs[0].add_run())

        header_paragraph = doc.sections[0].header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        static_header_text_run = header_paragraph.add_run()
        static_header_text_run.text = f"LIFT Version {get_version()}\t{datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"

        word_stream = io.BytesIO()
        doc.save(word_stream)
        word_stream.seek(0)
        word_bytes = word_stream.read()

        return word_bytes
